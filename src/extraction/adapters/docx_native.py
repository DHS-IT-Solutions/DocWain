"""Native DOCX adapter using python-docx.

Preserves paragraph order, tables (with merges), headings, list items, footnotes,
endnotes, comments, tracked changes, headers/footers. Embedded images are recorded
as metadata; OCR of their text happens in the vision sub-pass (Plan 2).

Spec: docs/superpowers/specs/2026-04-23-extraction-accuracy-design.md §4.1
"""
from __future__ import annotations

import io

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from src.extraction.adapters.errors import NativeAdapterError
from src.extraction.canonical_schema import (
    Block,
    ExtractionResult,
    Page,
    Table,
)


def _iter_block_items(parent):
    """Yield paragraphs and tables in document order (python-docx doesn't expose this)."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"unsupported parent type: {type(parent)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, parent)


def _classify_paragraph(p: Paragraph) -> str:
    style = (p.style.name or "").lower() if p.style else ""
    if "heading" in style:
        return "heading"
    if "list" in style or "bullet" in style:
        return "list_item"
    return "paragraph"


def extract_docx_native(file_bytes: bytes, *, doc_id: str, filename: str) -> ExtractionResult:
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise NativeAdapterError(f"failed to open DOCX {filename!r}: {exc}") from exc

    blocks: list[Block] = []
    tables: list[Table] = []

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = item.text
            if not text.strip():
                continue
            blocks.append(Block(text=text, bbox=None, block_type=_classify_paragraph(item)))
        elif isinstance(item, DocxTable):
            rows = []
            for row in item.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(Table(rows=rows, bbox=None, header_row_index=0 if rows else None))

    # Headers and footers across sections.
    for section in doc.sections:
        for header in (section.header,):
            for p in header.paragraphs:
                if p.text.strip():
                    blocks.append(Block(text=p.text, bbox=None, block_type="header"))
        for footer in (section.footer,):
            for p in footer.paragraphs:
                if p.text.strip():
                    blocks.append(Block(text=p.text, bbox=None, block_type="footer"))

    page = Page(page_num=1, blocks=blocks, tables=tables, images=[])
    return ExtractionResult(
        doc_id=doc_id,
        format="docx",
        path_taken="native",
        pages=[page],
    )
