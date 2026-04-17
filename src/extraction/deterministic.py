"""Deterministic extraction — faithful content capture without AI.

Layer 1 of the extraction pipeline. Given raw file bytes, returns a
``RawExtraction`` containing every piece of content the file actually
carries: text, tables, headers/footers, watermarks, image OCR, metadata.
No interpretation, no LLM, no domain assumptions.

Each format is handled by a dedicated parser:

- PDF   : pdfplumber (text + tables) with positional-word fallback for
          rule-less tables and cross-page table stitching
- DOCX  : python-docx + raw-XML watermark scan; table cell text is
          included in ``text_full``
- XLSX  : openpyxl with merged-cell flattening and formula resolution
- IMAGE : PIL preprocessing (grayscale, auto-contrast) + Tesseract with
          per-line output preserved
- CSV / TSV : pandas
- TXT / MD  : direct read

The output shape is stable regardless of source format so downstream
callers (merger, V2 interpreter, chunker) consume one API.

Error handling: the function never raises. Unrecoverable problems are
reported as ``warnings`` on the returned ``RawExtraction``, with content
captured up to the failure point.
"""

from __future__ import annotations

import io
import logging
import os
import zipfile
from dataclasses import dataclass, field
from enum import Enum
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public types
# ---------------------------------------------------------------------------


class BlockType(str, Enum):
    PARAGRAPH = "paragraph"
    HEADING = "heading"
    TABLE = "table"
    LIST_ITEM = "list_item"
    HEADER = "header"  # page / section header band
    FOOTER = "footer"
    WATERMARK = "watermark"
    IMAGE_TEXT = "image_text"  # OCR'd text line from an image
    CELL = "cell"  # individual XLSX cell
    FORMULA = "formula"
    CAPTION = "caption"


@dataclass
class Block:
    """A single unit of extracted content with provenance."""

    type: BlockType
    text: str
    page: Optional[int] = None
    bbox: Optional[Tuple[float, float, float, float]] = None  # (x0, y0, x1, y1)
    source_method: str = ""  # e.g. "pdfplumber_text", "tesseract_line"
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Table:
    """A structured table. ``rows`` includes the header as row[0] when known."""

    rows: List[List[str]]
    headers: List[str] = field(default_factory=list)
    page: Optional[int] = None
    bbox: Optional[Tuple[float, float, float, float]] = None
    source_method: str = ""
    cross_page: bool = False  # True when reconstructed from multi-page parts
    metadata: Dict[str, Any] = field(default_factory=dict)

    @property
    def n_rows(self) -> int:
        return len(self.rows)

    @property
    def n_cols(self) -> int:
        return max((len(r) for r in self.rows), default=0)


@dataclass
class RawExtraction:
    """Deterministic extraction output. One instance per input file."""

    file_format: str
    file_size_bytes: int
    filename: str
    text_full: str  # every readable piece of content, in reading order
    blocks: List[Block] = field(default_factory=list)
    tables: List[Table] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "file_format": self.file_format,
            "file_size_bytes": self.file_size_bytes,
            "filename": self.filename,
            "text_char_count": len(self.text_full),
            "text_full": self.text_full,
            "block_count": len(self.blocks),
            "blocks": [
                {
                    "type": b.type.value,
                    "text": b.text,
                    "page": b.page,
                    "bbox": b.bbox,
                    "source_method": b.source_method,
                    "metadata": b.metadata,
                }
                for b in self.blocks
            ],
            "table_count": len(self.tables),
            "tables": [
                {
                    "rows": t.rows,
                    "headers": t.headers,
                    "n_rows": t.n_rows,
                    "n_cols": t.n_cols,
                    "page": t.page,
                    "bbox": t.bbox,
                    "source_method": t.source_method,
                    "cross_page": t.cross_page,
                    "metadata": t.metadata,
                }
                for t in self.tables
            ],
            "metadata": self.metadata,
            "warnings": self.warnings,
        }


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".xls",
    ".png", ".jpg", ".jpeg", ".tiff", ".bmp",
    ".csv", ".tsv", ".txt", ".md",
}


def extract(content: bytes, filename: str) -> RawExtraction:
    """Route to the appropriate format handler.

    Returns an empty-but-valid ``RawExtraction`` with a warning if the
    extension is unsupported or the parser fails unrecoverably.
    """
    ext = os.path.splitext(filename)[1].lower()
    base = RawExtraction(
        file_format=ext.lstrip(".") or "unknown",
        file_size_bytes=len(content),
        filename=os.path.basename(filename),
        text_full="",
    )

    if ext not in SUPPORTED_EXTENSIONS:
        base.warnings.append(f"Unsupported extension: {ext}")
        return base

    try:
        if ext == ".pdf":
            return _extract_pdf(content, filename)
        if ext == ".docx":
            return _extract_docx(content, filename)
        if ext in (".xlsx", ".xls"):
            return _extract_xlsx(content, filename)
        if ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
            return _extract_image(content, filename)
        if ext in (".csv", ".tsv"):
            return _extract_csv(content, filename)
        if ext in (".txt", ".md"):
            return _extract_text(content, filename)
    except Exception as exc:  # noqa: BLE001
        logger.exception("Deterministic extraction failed for %s", filename)
        base.warnings.append(f"{type(exc).__name__}: {exc}")

    return base


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _extract_pdf(content: bytes, filename: str) -> RawExtraction:
    import fitz  # PyMuPDF
    import pdfplumber

    raw = RawExtraction(
        file_format="pdf",
        file_size_bytes=len(content),
        filename=os.path.basename(filename),
        text_full="",
    )

    # 1) Metadata via fitz
    try:
        with fitz.open(stream=content, filetype="pdf") as doc:
            raw.metadata["pdf_metadata"] = dict(doc.metadata or {})
            raw.metadata["page_count"] = doc.page_count
            raw.metadata["is_encrypted"] = bool(doc.is_encrypted)
            image_count = 0
            for page in doc:
                image_count += len(page.get_images(full=True))
            raw.metadata["embedded_image_count"] = image_count
    except Exception as exc:  # noqa: BLE001
        raw.warnings.append(f"fitz metadata failed: {exc}")

    # 2) Text + tables via pdfplumber
    page_texts: List[str] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                page_texts.append(page_text)
                if page_text.strip():
                    raw.blocks.append(Block(
                        type=BlockType.PARAGRAPH,
                        text=page_text,
                        page=page_idx,
                        source_method="pdfplumber_text",
                    ))

                # Default lattice/lines table strategy
                page_tables = page.extract_tables() or []
                # Fallback: try text-positional strategy ONLY when lattice
                # found nothing. A small table from lattice is usually a
                # genuine continuation on the page; don't displace it with a
                # coarser text-strategy table that may capture surrounding
                # non-tabular text (totals section, notes, etc.).
                if not page_tables:
                    try:
                        text_tables = page.extract_tables({
                            "vertical_strategy": "text",
                            "horizontal_strategy": "text",
                        }) or []
                        if text_tables:
                            page_tables = text_tables
                    except Exception as exc:  # noqa: BLE001
                        raw.warnings.append(f"text-strategy tables failed p{page_idx}: {exc}")

                for t_idx, table_grid in enumerate(page_tables, start=1):
                    if not table_grid:
                        continue
                    normalised = [
                        [("" if cell is None else str(cell).strip()) for cell in row]
                        for row in table_grid
                    ]
                    headers = normalised[0] if normalised else []
                    raw.tables.append(Table(
                        rows=normalised,
                        headers=headers,
                        page=page_idx,
                        source_method="pdfplumber_tables",
                        metadata={"table_index_on_page": t_idx},
                    ))
    except Exception as exc:  # noqa: BLE001
        raw.warnings.append(f"pdfplumber failed: {exc}")

    # 3) Merge adjacent-page tables when page N+1 starts with no header that
    #    matches the header of page N's last table and has the same col count
    #    (the "split tables across pages" case).
    raw.tables = _stitch_cross_page_tables(raw.tables)

    raw.text_full = "\n\n".join(t for t in page_texts if t.strip())
    return raw


def _stitch_cross_page_tables(tables: List[Table]) -> List[Table]:
    """Merge consecutive tables that look like a single table split by a page break.

    Heuristic: for tables T_k on page P and T_{k+1} on page P+1, stitch if:
    - same n_cols
    - T_{k+1}'s first row does NOT match T_k's headers (i.e., page 2 has no repeated header)
    - OR T_{k+1}'s first row IS the same as T_k's headers (repeated header — drop it and stitch)
    """
    if len(tables) < 2:
        return tables

    stitched: List[Table] = []
    i = 0
    while i < len(tables):
        current = tables[i]
        j = i + 1
        while j < len(tables):
            nxt = tables[j]
            same_cols = current.n_cols == nxt.n_cols and current.n_cols > 0
            adjacent = (current.page is not None and nxt.page is not None
                        and nxt.page == current.page + (j - i))
            if not (same_cols and adjacent):
                break

            first_row = nxt.rows[0] if nxt.rows else []
            header_repeat = (
                bool(current.headers)
                and first_row
                and all(
                    (str(first_row[k]).strip().lower()
                     == str(current.headers[k]).strip().lower())
                    for k in range(min(len(first_row), len(current.headers)))
                )
            )
            body = nxt.rows[1:] if header_repeat else list(nxt.rows)
            current = Table(
                rows=current.rows + body,
                headers=current.headers,
                page=current.page,
                source_method=f"{current.source_method}+stitched",
                cross_page=True,
                metadata={**current.metadata, "stitched_from_pages": [current.page, nxt.page]},
            )
            j += 1
        stitched.append(current)
        i = j
    return stitched


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _extract_docx(content: bytes, filename: str) -> RawExtraction:
    from docx import Document

    raw = RawExtraction(
        file_format="docx",
        file_size_bytes=len(content),
        filename=os.path.basename(filename),
        text_full="",
    )

    doc = Document(io.BytesIO(content))

    # Paragraphs
    all_text_parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        btype = BlockType.HEADING if (p.style and "Heading" in (p.style.name or "")) else BlockType.PARAGRAPH
        raw.blocks.append(Block(type=btype, text=t, source_method="python-docx_paragraphs"))
        all_text_parts.append(t)

    # Tables — include cell text in text_full so "content-only-in-tables" docs
    # don't look empty downstream.
    for t_idx, table in enumerate(doc.tables, start=1):
        grid: List[List[str]] = []
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            grid.append(cells)
            all_text_parts.extend(c for c in cells if c)
        if not grid:
            continue
        headers = grid[0] if grid else []
        raw.tables.append(Table(
            rows=grid,
            headers=headers,
            source_method="python-docx_tables",
            metadata={"table_index": t_idx},
        ))

    # Headers / footers
    header_lines: List[str] = []
    footer_lines: List[str] = []
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr is None:
                continue
            for p in hdr.paragraphs:
                t = (p.text or "").strip()
                if t and t not in header_lines:
                    header_lines.append(t)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr is None:
                continue
            for p in ftr.paragraphs:
                t = (p.text or "").strip()
                if t and t not in footer_lines:
                    footer_lines.append(t)
    for t in header_lines:
        raw.blocks.append(Block(type=BlockType.HEADER, text=t, source_method="python-docx_section_headers"))
    for t in footer_lines:
        raw.blocks.append(Block(type=BlockType.FOOTER, text=t, source_method="python-docx_section_footers"))
    raw.metadata["headers"] = header_lines
    raw.metadata["footers"] = footer_lines

    # Watermarks — scan header XML for v:textpath elements (Word's watermark shape)
    watermarks: List[str] = []
    try:
        import re as _re
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            for name in z.namelist():
                if "header" in name and name.endswith(".xml"):
                    xml = z.read(name).decode("utf-8", errors="ignore")
                    for m in _re.finditer(r'string="([^"]+)"', xml):
                        candidate = m.group(1).strip()
                        if candidate and candidate not in watermarks:
                            watermarks.append(candidate)
    except Exception as exc:  # noqa: BLE001
        raw.warnings.append(f"watermark scan failed: {exc}")
    for w in watermarks:
        raw.blocks.append(Block(type=BlockType.WATERMARK, text=w, source_method="docx_xml_textpath"))
    raw.metadata["watermarks"] = watermarks

    # Inline image count
    try:
        inline_images = len([r for r in doc.part.related_parts.values()
                             if r.content_type.startswith("image/")])
    except Exception:  # noqa: BLE001
        inline_images = -1
    raw.metadata["inline_image_count"] = inline_images
    raw.metadata["paragraph_count"] = sum(1 for p in doc.paragraphs if (p.text or "").strip())

    raw.text_full = "\n".join(all_text_parts)
    return raw


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def _extract_xlsx(content: bytes, filename: str) -> RawExtraction:
    import openpyxl

    raw = RawExtraction(
        file_format="xlsx",
        file_size_bytes=len(content),
        filename=os.path.basename(filename),
        text_full="",
    )

    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    # Load formulas too — keep alongside resolved values for provenance.
    try:
        wb_formulas = openpyxl.load_workbook(io.BytesIO(content), data_only=False)
    except Exception:  # noqa: BLE001
        wb_formulas = None

    raw.metadata["sheet_names"] = wb.sheetnames
    raw.metadata["sheet_count"] = len(wb.sheetnames)

    text_parts: List[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        max_row = ws.max_row or 0
        max_col = ws.max_column or 0

        # Build a dense grid of values, then flatten merged ranges so every
        # cell in a merged block carries the top-left value.
        grid: List[List[str]] = [
            ["" for _ in range(max_col)] for _ in range(max_row)
        ]
        for row in ws.iter_rows(min_row=1, max_row=max_row, values_only=False):
            for cell in row:
                if cell.row and cell.column:
                    val = cell.value
                    grid[cell.row - 1][cell.column - 1] = (
                        "" if val is None else str(val)
                    )

        merged_ranges = list(ws.merged_cells.ranges)
        for rng in merged_ranges:
            top_left_value = ws.cell(row=rng.min_row, column=rng.min_col).value
            v = "" if top_left_value is None else str(top_left_value)
            for r in range(rng.min_row, rng.max_row + 1):
                for c in range(rng.min_col, rng.max_col + 1):
                    grid[r - 1][c - 1] = v

        # Collect cell text blocks + sheet as a Table
        non_empty_count = 0
        for r_idx, row in enumerate(grid, start=1):
            for c_idx, val in enumerate(row, start=1):
                if val:
                    non_empty_count += 1
                    raw.blocks.append(Block(
                        type=BlockType.CELL,
                        text=val,
                        source_method="openpyxl_cell",
                        metadata={
                            "sheet": sheet_name,
                            "row": r_idx,
                            "col": c_idx,
                        },
                    ))
                    text_parts.append(val)

        raw.tables.append(Table(
            rows=grid,
            headers=grid[0] if grid else [],
            source_method="openpyxl_sheet",
            metadata={
                "sheet": sheet_name,
                "dimensions": f"{max_row}r x {max_col}c",
                "non_empty_cells": non_empty_count,
                "merged_cell_ranges": [str(r) for r in merged_ranges],
            },
        ))

        # Formulas — attach as dedicated blocks
        if wb_formulas is not None:
            try:
                ws_f = wb_formulas[sheet_name]
                for row in ws_f.iter_rows(min_row=1, max_row=max_row, values_only=False):
                    for cell in row:
                        if isinstance(cell.value, str) and cell.value.startswith("="):
                            raw.blocks.append(Block(
                                type=BlockType.FORMULA,
                                text=cell.value,
                                source_method="openpyxl_formula",
                                metadata={
                                    "sheet": sheet_name,
                                    "row": cell.row,
                                    "col": cell.column,
                                },
                            ))
            except Exception as exc:  # noqa: BLE001
                raw.warnings.append(f"formula scan failed sheet={sheet_name}: {exc}")

    raw.text_full = "\n".join(text_parts)
    return raw


# ---------------------------------------------------------------------------
# IMAGE
# ---------------------------------------------------------------------------


def _extract_image(content: bytes, filename: str) -> RawExtraction:
    from PIL import Image, ImageOps
    import pytesseract

    raw = RawExtraction(
        file_format=os.path.splitext(filename)[1].lstrip(".").lower() or "image",
        file_size_bytes=len(content),
        filename=os.path.basename(filename),
        text_full="",
    )

    try:
        img = Image.open(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raw.warnings.append(f"PIL open failed: {exc}")
        return raw

    raw.metadata["pixel_dimensions"] = f"{img.width}x{img.height}"
    raw.metadata["mode"] = img.mode

    # Preprocess: grayscale + auto-contrast. Keep it conservative; aggressive
    # thresholding/deskew can make Tesseract worse on clean scans.
    processed = img.convert("L")
    try:
        processed = ImageOps.autocontrast(processed)
    except Exception as exc:  # noqa: BLE001
        raw.warnings.append(f"autocontrast failed: {exc}")

    # Flat text
    try:
        flat_text = pytesseract.image_to_string(processed)
    except Exception as exc:  # noqa: BLE001
        raw.warnings.append(f"tesseract image_to_string failed: {exc}")
        flat_text = ""

    # Per-line text via image_to_data — preserves line structure.
    try:
        data = pytesseract.image_to_data(processed, output_type=pytesseract.Output.DICT)
    except Exception as exc:  # noqa: BLE001
        raw.warnings.append(f"tesseract image_to_data failed: {exc}")
        data = None

    lines: List[str] = []
    confidences: List[int] = []
    if data:
        # Group by (page_num, block_num, par_num, line_num) — tesseract keys.
        line_buckets: Dict[Tuple[int, int, int, int], List[Tuple[int, str, int]]] = {}
        for i, word in enumerate(data.get("text", [])):
            if not word or not word.strip():
                continue
            key = (
                data["page_num"][i],
                data["block_num"][i],
                data["par_num"][i],
                data["line_num"][i],
            )
            x_left = int(data["left"][i])
            y_top = int(data["top"][i])
            conf_raw = data["conf"][i]
            try:
                conf = int(conf_raw)
            except (TypeError, ValueError):
                conf = -1
            if conf >= 0:
                confidences.append(conf)
            line_buckets.setdefault(key, []).append((x_left, word, y_top))

        for key, words in sorted(
            line_buckets.items(),
            key=lambda kv: (kv[0], min((w[2] for w in kv[1]), default=0)),
        ):
            words_sorted = sorted(words, key=lambda w: w[0])
            line_text = " ".join(w[1] for w in words_sorted)
            if not line_text.strip():
                continue
            lines.append(line_text)
            raw.blocks.append(Block(
                type=BlockType.IMAGE_TEXT,
                text=line_text,
                source_method="tesseract_line",
                metadata={
                    "block_num": key[1],
                    "par_num": key[2],
                    "line_num": key[3],
                    "min_x": min(w[0] for w in words_sorted),
                },
            ))

    if confidences:
        raw.metadata["ocr_confidence_mean"] = round(sum(confidences) / len(confidences), 1)
        raw.metadata["ocr_confidence_min"] = min(confidences)
        raw.metadata["ocr_confidence_max"] = max(confidences)
    raw.metadata["ocr_word_count"] = sum(1 for w in (data.get("text", []) if data else []) if w and w.strip())
    raw.metadata["ocr_line_count"] = len(lines)

    # Prefer structured line-joined text when available; fall back to flat.
    raw.text_full = "\n".join(lines) if lines else flat_text

    return raw


# ---------------------------------------------------------------------------
# CSV / TSV
# ---------------------------------------------------------------------------


def _extract_csv(content: bytes, filename: str) -> RawExtraction:
    import pandas as pd

    ext = os.path.splitext(filename)[1].lower()
    sep = "\t" if ext == ".tsv" else ","
    raw = RawExtraction(
        file_format=ext.lstrip("."),
        file_size_bytes=len(content),
        filename=os.path.basename(filename),
        text_full="",
    )
    df = pd.read_csv(io.BytesIO(content), sep=sep, low_memory=False)
    raw.metadata["row_count"] = len(df)
    raw.metadata["col_count"] = len(df.columns)
    raw.metadata["columns"] = list(df.columns)

    rows = [[str(c) for c in df.columns]]
    rows.extend([[str(v) for v in row] for row in df.itertuples(index=False, name=None)])
    raw.tables.append(Table(
        rows=rows,
        headers=list(df.columns),
        source_method="pandas_csv",
    ))

    text_parts: List[str] = []
    for row in rows:
        text_parts.append(" | ".join(row))
    raw.text_full = "\n".join(text_parts)
    return raw


# ---------------------------------------------------------------------------
# TXT / MD
# ---------------------------------------------------------------------------


def _extract_text(content: bytes, filename: str) -> RawExtraction:
    raw = RawExtraction(
        file_format=os.path.splitext(filename)[1].lstrip(".") or "txt",
        file_size_bytes=len(content),
        filename=os.path.basename(filename),
        text_full="",
    )
    text = content.decode("utf-8", errors="replace")
    raw.text_full = text
    raw.metadata["line_count"] = text.count("\n") + 1
    for para in text.split("\n\n"):
        p = para.strip()
        if p:
            raw.blocks.append(Block(
                type=BlockType.PARAGRAPH,
                text=p,
                source_method="plain_text",
            ))
    return raw


__all__ = [
    "Block",
    "BlockType",
    "Table",
    "RawExtraction",
    "SUPPORTED_EXTENSIONS",
    "extract",
]
