import io

from docx import Document

from src.extraction.adapters.docx_native import extract_docx_native
from src.extraction.canonical_schema import ExtractionResult


def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, cell_text in enumerate(row):
                t.rows[i].cells[j].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_native_extracts_paragraphs():
    docx = _make_docx(["first paragraph", "second paragraph", "third"])
    result = extract_docx_native(docx, doc_id="d1", filename="t.docx")
    assert isinstance(result, ExtractionResult)
    assert result.format == "docx"
    assert result.path_taken == "native"
    assert len(result.pages) == 1
    texts = [b.text for b in result.pages[0].blocks]
    assert "first paragraph" in texts
    assert "second paragraph" in texts
    assert "third" in texts


def test_docx_native_extracts_tables():
    docx = _make_docx(["before table"], table_rows=[["h1", "h2"], ["r1c1", "r1c2"]])
    result = extract_docx_native(docx, doc_id="d2", filename="t.docx")
    assert len(result.pages[0].tables) == 1
    assert result.pages[0].tables[0].rows == [["h1", "h2"], ["r1c1", "r1c2"]]
